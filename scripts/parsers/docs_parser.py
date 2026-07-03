import logging
from pathlib import Path
from typing import List

# 프로젝트의 기본 베이스 클래스와 결과 데이터 구조를 임포트합니다.
from parsers.base import BaseParser, ParsedContent

logger = logging.getLogger(__name__)

class DOCXParser(BaseParser):
    """python-docx 라이브러리를 기반으로 DOCX 문서의 텍스트와 표를 추출하는 클래스입니다."""

    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]

    def parse(self, file_path: Path, **kwargs) -> ParsedContent:
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx 라이브러리가 설치되지 않았습니다. 'pip install python-docx'를 실행하세요.")
            return ParsedContent(text="", parse_method="docx_unavailable", confidence=0.0)

        file_path = Path(file_path)
        if not file_path.exists():
            return ParsedContent(text="", parse_method="file_not_found", confidence=0.0)

        try:
            # DOCX 파일 로드
            doc = Document(str(file_path))
        except Exception as e:
            logger.error(f"DOCX 파일을 여는 중 오류 발생: {file_path} -> {e}")
            return ParsedContent(text="", parse_method="open_failed", confidence=0.0)

        paragraphs_text: List[str] = []
        tables_markdown: List[str] = []

        # 1. 문단(Paragraph) 추출
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs_text.append(text)

        # 2. 표(Table) 추출 및 마크다운 변환
        for i, table in enumerate(doc.tables, 1):
            md = self._table_to_markdown(table)
            if md:
                tables_markdown.append(f"### 문서 내 표 {i}\n{md}")

        # 추출된 텍스트 결합
        full_text = "\n\n".join(paragraphs_text)

        return ParsedContent(
            text=full_text,
            tables=tables_markdown,
            metadata={
                "filename": file_path.name,
                "paragraph_count": len(paragraphs_text),
                "table_count": len(doc.tables),
                "file_size": file_path.stat().st_size, # 파일 크기 정보 포함[cite: 1]
            },
            parse_method="python-docx",
            confidence=1.0 if paragraphs_text else 0.3 # 텍스트가 없으면 신뢰도를 낮춤[cite: 1]
        )

    def _table_to_markdown(self, table) -> str:
        """python-docx 표 객체를 마크다운 문자열로 변환합니다[cite: 1]."""
        try:
            rows = []
            for i, row in enumerate(table.rows):
                # 셀 내부 줄바꿈 제거 및 공백 정리
                cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
                
                # 헤더 구분선 추가
                if i == 0:
                    rows.append("|" + "|".join(["---" for _ in row.cells]) + "|")
            
            return "\n".join(rows)
        except Exception as e:
            logger.warning(f"표 변환 중 오류 발생: {e}")
            return ""